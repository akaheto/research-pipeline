"""Tests for report generators."""

import pytest
import tempfile
from pathlib import Path
from datetime import datetime

from research_pipeline.models import Finding, Source, Research
from research_pipeline.reporters.markdown_reporter import MarkdownReporter
from research_pipeline.reporters.docx_reporter import DocxReporter
from research_pipeline.reporters.orchestrator import ReportingOrchestrator


@pytest.fixture
def sample_research():
    """Create sample research for testing."""
    findings = [
        Finding(
            text="AI is transforming industries worldwide",
            source=Source(
                url="https://example1.com",
                title="Tech News Daily",
                domain="example1.com"
            ),
            relevance_score=0.95,
            credibility_score=0.85,
            recency_score=0.9
        ),
        Finding(
            text="Machine learning adoption accelerating",
            source=Source(
                url="https://example2.com",
                title="AI Weekly",
                domain="example2.com"
            ),
            relevance_score=0.88,
            credibility_score=0.92,
            recency_score=0.85
        ),
    ]
    return Research(
        topic="AI trends 2024",
        findings=findings,
        knowledge_base_insights=["Previously tracked ML adoption rates", "Competitor analysis from Q3"],
        summary="AI transformation is accelerating across all sectors"
    )


class TestMarkdownReporter:
    """Tests for Markdown report generator."""

    @pytest.mark.asyncio
    async def test_generates_markdown_file(self, sample_research):
        """Test that Markdown report is generated."""
        with tempfile.TemporaryDirectory() as tmpdir:
            reporter = MarkdownReporter()
            output_path = Path(tmpdir) / "report.md"

            result = await reporter.generate(sample_research, output_path)

            assert result.exists()
            assert result.suffix == ".md"
            content = result.read_text()
            assert "AI trends 2024" in content

    @pytest.mark.asyncio
    async def test_markdown_contains_sections(self, sample_research):
        """Test that Markdown includes all report sections."""
        with tempfile.TemporaryDirectory() as tmpdir:
            reporter = MarkdownReporter()
            output_path = Path(tmpdir) / "report.md"

            await reporter.generate(sample_research, output_path)
            content = output_path.read_text()

            assert "# Research Report:" in content
            assert "## Summary" in content
            assert "## Key Findings" in content
            assert "## Related Knowledge Base Insights" in content
            assert "## All Sources" in content

    @pytest.mark.asyncio
    async def test_markdown_contains_findings(self, sample_research):
        """Test that Markdown includes finding details."""
        with tempfile.TemporaryDirectory() as tmpdir:
            reporter = MarkdownReporter()
            output_path = Path(tmpdir) / "report.md"

            await reporter.generate(sample_research, output_path)
            content = output_path.read_text()

            assert "AI is transforming industries" in content
            assert "Tech News Daily" in content
            assert "example1.com" in content
            assert "95.0%" in content  # Relevance score formatted

    @pytest.mark.asyncio
    async def test_markdown_contains_kb_insights(self, sample_research):
        """Test that Markdown includes knowledge base insights."""
        with tempfile.TemporaryDirectory() as tmpdir:
            reporter = MarkdownReporter()
            output_path = Path(tmpdir) / "report.md"

            await reporter.generate(sample_research, output_path)
            content = output_path.read_text()

            assert "Previously tracked ML adoption rates" in content
            assert "Competitor analysis from Q3" in content

    @pytest.mark.asyncio
    async def test_markdown_empty_findings(self):
        """Test Markdown with no findings."""
        with tempfile.TemporaryDirectory() as tmpdir:
            reporter = MarkdownReporter()
            research = Research(topic="Empty test", findings=[])
            output_path = Path(tmpdir) / "report.md"

            result = await reporter.generate(research, output_path)
            content = result.read_text()

            assert "Empty test" in content
            assert "Key Findings" not in content


class TestDocxReporter:
    """Tests for Word document reporter."""

    @pytest.mark.asyncio
    async def test_generates_docx_file(self, sample_research):
        """Test that Word report is generated."""
        with tempfile.TemporaryDirectory() as tmpdir:
            reporter = DocxReporter()
            output_path = Path(tmpdir) / "report.docx"

            result = await reporter.generate(sample_research, output_path)

            assert result.exists()
            assert result.suffix == ".docx"

    @pytest.mark.asyncio
    async def test_docx_contains_title(self, sample_research):
        """Test that Word document includes title."""
        with tempfile.TemporaryDirectory() as tmpdir:
            reporter = DocxReporter()
            output_path = Path(tmpdir) / "report.docx"

            await reporter.generate(sample_research, output_path)

            from docx import Document
            doc = Document(output_path)
            heading_text = doc.paragraphs[0].text

            assert "AI trends 2024" in heading_text

    @pytest.mark.asyncio
    async def test_docx_contains_findings(self, sample_research):
        """Test that Word document includes findings."""
        with tempfile.TemporaryDirectory() as tmpdir:
            reporter = DocxReporter()
            output_path = Path(tmpdir) / "report.docx"

            await reporter.generate(sample_research, output_path)

            from docx import Document
            doc = Document(output_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])

            assert "AI is transforming industries" in full_text
            assert "Tech News Daily" in full_text

    @pytest.mark.asyncio
    async def test_docx_contains_tables(self, sample_research):
        """Test that Word document includes finding tables."""
        with tempfile.TemporaryDirectory() as tmpdir:
            reporter = DocxReporter()
            output_path = Path(tmpdir) / "report.docx"

            await reporter.generate(sample_research, output_path)

            from docx import Document
            doc = Document(output_path)

            # Should have tables for findings
            assert len(doc.tables) > 0

    @pytest.mark.asyncio
    async def test_docx_empty_findings(self):
        """Test Word document with no findings."""
        with tempfile.TemporaryDirectory() as tmpdir:
            reporter = DocxReporter()
            research = Research(topic="Empty test", findings=[])
            output_path = Path(tmpdir) / "report.docx"

            result = await reporter.generate(research, output_path)

            from docx import Document
            doc = Document(result)
            full_text = "\n".join([p.text for p in doc.paragraphs])

            assert "Empty test" in full_text


class TestReportingOrchestrator:
    """Tests for reporting orchestrator."""

    @pytest.mark.asyncio
    async def test_generate_markdown(self, sample_research):
        """Test generating Markdown report."""
        with tempfile.TemporaryDirectory() as tmpdir:
            orchestrator = ReportingOrchestrator()
            output_path = Path(tmpdir) / "report"

            result = await orchestrator.generate(
                sample_research,
                str(output_path),
                format="markdown"
            )

            assert result.suffix == ".md"
            assert result.exists()

    @pytest.mark.asyncio
    async def test_generate_docx(self, sample_research):
        """Test generating Word report."""
        with tempfile.TemporaryDirectory() as tmpdir:
            orchestrator = ReportingOrchestrator()
            output_path = Path(tmpdir) / "report"

            result = await orchestrator.generate(
                sample_research,
                str(output_path),
                format="docx"
            )

            assert result.suffix == ".docx"
            assert result.exists()

    @pytest.mark.asyncio
    async def test_generate_with_correct_extension(self, sample_research):
        """Test that correct extension is enforced."""
        with tempfile.TemporaryDirectory() as tmpdir:
            orchestrator = ReportingOrchestrator()

            # Request .md but with wrong extension
            result = await orchestrator.generate(
                sample_research,
                str(Path(tmpdir) / "report.docx"),
                format="markdown"
            )

            # Should correct to .md
            assert result.suffix == ".md"

    @pytest.mark.asyncio
    async def test_unsupported_format(self, sample_research):
        """Test that unsupported format raises error."""
        with tempfile.TemporaryDirectory() as tmpdir:
            orchestrator = ReportingOrchestrator()

            with pytest.raises(ValueError, match="Unsupported format"):
                await orchestrator.generate(
                    sample_research,
                    str(Path(tmpdir) / "report"),
                    format="pdf"
                )

    @pytest.mark.asyncio
    async def test_generate_both_formats(self, sample_research):
        """Test generating reports in both formats."""
        with tempfile.TemporaryDirectory() as tmpdir:
            orchestrator = ReportingOrchestrator()

            results = await orchestrator.generate_both(
                sample_research,
                output_dir=tmpdir
            )

            assert "markdown" in results
            assert "docx" in results
            assert results["markdown"].exists()
            assert results["docx"].exists()
            assert results["markdown"].suffix == ".md"
            assert results["docx"].suffix == ".docx"

    @pytest.mark.asyncio
    async def test_generate_both_creates_directory(self, sample_research):
        """Test that generate_both creates output directory."""
        with tempfile.TemporaryDirectory() as tmpdir:
            orchestrator = ReportingOrchestrator()
            new_dir = Path(tmpdir) / "new_output_dir"

            results = await orchestrator.generate_both(
                sample_research,
                output_dir=str(new_dir)
            )

            assert new_dir.exists()
            assert results["markdown"].exists()
            assert results["docx"].exists()
