"""End-to-end integration tests for the research pipeline."""

import pytest
import tempfile
from pathlib import Path
from datetime import datetime
from unittest.mock import AsyncMock, MagicMock, patch

from research_pipeline.models import Finding, Source, Research
from research_pipeline.collectors.orchestrator import DataCollectionOrchestrator
from research_pipeline.processors.orchestrator import ProcessingOrchestrator
from research_pipeline.reporters.orchestrator import ReportingOrchestrator


@pytest.fixture
def realistic_findings():
    """Create realistic research findings."""
    return [
        Finding(
            text="AI adoption accelerating across enterprise sectors with 73% planning investments",
            source=Source(
                url="https://techinsights.example.com/ai-adoption-2024",
                title="Tech Insights: AI Investment Trends",
                domain="techinsights.example.com"
            ),
            relevance_score=0.95,
            credibility_score=0.92,
            recency_score=0.98
        ),
        Finding(
            text="Machine learning models achieving state-of-the-art performance in NLP tasks",
            source=Source(
                url="https://arxiv.example.com/paper-2024-001",
                title="Recent Advances in NLP",
                domain="arxiv.example.com"
            ),
            relevance_score=0.88,
            credibility_score=0.95,
            recency_score=0.92
        ),
        Finding(
            text="Large language models now handling multimodal inputs more effectively",
            source=Source(
                url="https://airesearch.example.com/multimodal-2024",
                title="Multimodal AI Research Report",
                domain="airesearch.example.com"
            ),
            relevance_score=0.85,
            credibility_score=0.88,
            recency_score=0.95
        ),
        Finding(
            text="Cost of training large models continues to decrease with optimization techniques",
            source=Source(
                url="https://mlops.example.com/training-costs",
                title="MLOps Cost Analysis",
                domain="mlops.example.com"
            ),
            relevance_score=0.78,
            credibility_score=0.85,
            recency_score=0.88
        ),
    ]


class TestFullPipelineIntegration:
    """End-to-end pipeline tests."""

    @pytest.mark.asyncio
    async def test_full_pipeline_docx_output(self, realistic_findings):
        """Test complete pipeline from collection to Word report."""
        with tempfile.TemporaryDirectory() as tmpdir:
            research = Research(
                topic="AI trends 2024",
                findings=realistic_findings.copy(),
                knowledge_base_insights=["Prior Q3 analysis showed similar trends"]
            )

            # Process
            processor = ProcessingOrchestrator(max_results=5)
            research = await processor.process(research)

            # Generate report
            reporter = ReportingOrchestrator()
            output_path = Path(tmpdir) / "report.docx"
            result = await reporter.generate(research, output_path, format="docx")

            # Verify
            assert result.exists()
            assert result.suffix == ".docx"

            # Verify document content
            from docx import Document
            doc = Document(result)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            assert "AI trends 2024" in full_text
            assert len(doc.tables) > 0  # Should have finding tables

    @pytest.mark.asyncio
    async def test_full_pipeline_markdown_output(self, realistic_findings):
        """Test complete pipeline with Markdown output."""
        with tempfile.TemporaryDirectory() as tmpdir:
            research = Research(
                topic="Machine Learning",
                findings=realistic_findings.copy()
            )

            processor = ProcessingOrchestrator(max_results=3)
            research = await processor.process(research)

            reporter = ReportingOrchestrator()
            output_path = Path(tmpdir) / "report.md"
            result = await reporter.generate(research, output_path, format="markdown")

            # Verify
            assert result.exists()
            assert result.suffix == ".md"
            content = result.read_text()
            assert "Machine Learning" in content
            assert "##" in content  # Markdown headers

    @pytest.mark.asyncio
    async def test_pipeline_with_empty_findings(self):
        """Test pipeline gracefully handles empty findings."""
        with tempfile.TemporaryDirectory() as tmpdir:
            research = Research(
                topic="No results topic",
                findings=[]
            )

            processor = ProcessingOrchestrator()
            research = await processor.process(research)

            reporter = ReportingOrchestrator()
            result = await reporter.generate(research, Path(tmpdir) / "empty.docx")

            assert result.exists()

    @pytest.mark.asyncio
    async def test_pipeline_ranking_accuracy(self, realistic_findings):
        """Test that pipeline ranks findings by combined score correctly."""
        research = Research(
            topic="test",
            findings=realistic_findings.copy()
        )

        processor = ProcessingOrchestrator()
        result = await processor.process(research)

        # Verify rankings
        scores = [f.combined_score for f in result.findings]
        assert scores == sorted(scores, reverse=True)

        # First finding should be highest scoring
        assert result.findings[0].text == "AI adoption accelerating across enterprise sectors with 73% planning investments"

    @pytest.mark.asyncio
    async def test_pipeline_deduplication(self):
        """Test that pipeline removes duplicates correctly."""
        findings = [
            Finding(
                text="Machine learning is advancing",
                source=Source(url="https://a.com", title="A", domain="a.com"),
                relevance_score=0.9,
                credibility_score=0.8,
                recency_score=0.85
            ),
            Finding(
                text="Machine learning is advancing",  # Exact duplicate
                source=Source(url="https://b.com", title="B", domain="b.com"),
                relevance_score=0.8,
                credibility_score=0.85,
                recency_score=0.80
            ),
            Finding(
                text="Different content about something else",
                source=Source(url="https://c.com", title="C", domain="c.com"),
                relevance_score=0.5,
                credibility_score=0.6,
                recency_score=0.5
            ),
        ]

        research = Research(topic="test", findings=findings)
        processor = ProcessingOrchestrator(similarity_threshold=0.95)
        result = await processor.process(research)

        # Should have 2 findings (one duplicate removed)
        assert len(result.findings) == 2
        # Best-scoring duplicate should remain
        assert result.findings[0].relevance_score == 0.9


class TestErrorHandling:
    """Tests for error handling and edge cases."""

    @pytest.mark.asyncio
    async def test_processor_preserves_edge_case_scores(self):
        """Test processor handles edge case scores (0.0, 1.0)."""
        findings = [
            Finding(
                text="Test finding at score boundaries",
                source=Source(url="https://test.com", title="Test", domain="test.com"),
                relevance_score=0.0,  # Minimum
                credibility_score=1.0,  # Maximum
                recency_score=0.5
            )
        ]

        research = Research(topic="test", findings=findings)
        processor = ProcessingOrchestrator()
        result = await processor.process(research)

        # Edge case scores should be preserved
        assert len(result.findings) > 0
        assert result.findings[0].relevance_score == 0.0
        assert result.findings[0].credibility_score == 1.0

    @pytest.mark.asyncio
    async def test_reporter_creates_parent_directories(self):
        """Test that reporter creates missing parent directories."""
        with tempfile.TemporaryDirectory() as tmpdir:
            research = Research(topic="test", findings=[])
            reporter = ReportingOrchestrator()

            # Output path with non-existent parent directories
            nested_path = Path(tmpdir) / "deep" / "nested" / "dir" / "report.docx"

            result = await reporter.generate(research, nested_path)

            # Should have created all parent directories
            assert result.exists()
            assert result.parent.exists()

    @pytest.mark.asyncio
    async def test_max_results_limit(self):
        """Test that max_results limit is respected."""
        findings = [
            Finding(
                text=f"Unique finding about topic {i}: This discusses aspect number {i} in detail with specific information",
                source=Source(url=f"https://test{i}.com", title=f"Test {i}", domain=f"test{i}.com"),
                relevance_score=0.5 + (i * 0.003),  # Vary scores
                credibility_score=0.8,
                recency_score=0.7
            )
            for i in range(100)
        ]

        research = Research(topic="test", findings=findings)
        processor = ProcessingOrchestrator(max_results=10, similarity_threshold=0.99)
        result = await processor.process(research)

        # Should limit to 10
        assert len(result.findings) <= 10

    @pytest.mark.asyncio
    async def test_min_score_filtering(self):
        """Test that min_score filters out low-scoring findings."""
        findings = [
            Finding(
                text=f"Finding {i}",
                source=Source(url=f"https://test{i}.com", title=f"Test {i}", domain=f"test{i}.com"),
                relevance_score=0.1 + (i * 0.1),
                credibility_score=0.1 + (i * 0.1),
                recency_score=0.1 + (i * 0.1)
            )
            for i in range(10)
        ]

        research = Research(topic="test", findings=findings)
        processor = ProcessingOrchestrator(min_score=0.5)
        result = await processor.process(research)

        # All remaining findings should be above threshold
        for finding in result.findings:
            assert finding.combined_score >= 0.5


class TestPerformanceAndRobustness:
    """Tests for performance and robustness."""

    @pytest.mark.asyncio
    async def test_large_findings_set(self):
        """Test pipeline handles large number of findings."""
        findings = [
            Finding(
                text=f"Article {i}: Study shows significant impact in category {i % 10} with detailed analysis of trends and data",
                source=Source(
                    url=f"https://source{i}.example.com/article/{i}",
                    title=f"Source Document {i}",
                    domain=f"source{i}.example.com"
                ),
                relevance_score=0.5 + ((i % 100) * 0.005),
                credibility_score=0.6 + ((i % 100) * 0.004),
                recency_score=0.7 + ((i % 100) * 0.003)
            )
            for i in range(500)
        ]

        research = Research(topic="large dataset test", findings=findings)
        processor = ProcessingOrchestrator(max_results=50, similarity_threshold=0.99)

        # Should complete without error
        result = await processor.process(research)

        assert len(result.findings) > 0
        assert len(result.findings) <= 50
        assert result.summary is not None

    @pytest.mark.asyncio
    async def test_special_characters_in_content(self):
        """Test pipeline handles special characters."""
        with tempfile.TemporaryDirectory() as tmpdir:
            research = Research(
                topic="Special chars & symbols: café, naïve, 中文",
                findings=[
                    Finding(
                        text="Content with special chars: émojis 🤖 and symbols @#$%",
                        source=Source(
                            url="https://example.com/special",
                            title='Title with "quotes" and apostrophes',
                            domain="example.com"
                        ),
                        relevance_score=0.8,
                        credibility_score=0.8,
                        recency_score=0.8
                    )
                ],
                knowledge_base_insights=["Insight with ñ and unicode: 🔬"]
            )

            processor = ProcessingOrchestrator()
            research = await processor.process(research)

            reporter = ReportingOrchestrator()
            result = await reporter.generate(research, Path(tmpdir) / "special.docx")

            assert result.exists()

    @pytest.mark.asyncio
    async def test_very_long_finding_text(self):
        """Test pipeline handles very long finding text."""
        long_text = "This is a very long finding. " * 100  # ~3000 chars

        research = Research(
            topic="long content",
            findings=[
                Finding(
                    text=long_text,
                    source=Source(url="https://example.com", title="Long", domain="example.com"),
                    relevance_score=0.8,
                    credibility_score=0.8,
                    recency_score=0.8
                )
            ]
        )

        processor = ProcessingOrchestrator()
        result = await processor.process(research)

        # Should preserve full text
        assert len(result.findings[0].text) > 100
